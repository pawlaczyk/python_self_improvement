#coding=utf-8
import main2
import openpyxl
import docx
from openpyxl.styles import Font, Color
from openpyxl.styles import colors
from docx.shared import Pt
import uno
import tables
import os
import xml
import zipfile
from docx.shared import Cm
from docx.enum.style import WD_STYLE_TYPE
from docx.dml.color import ColorFormat

####################################33
#import rstwordml
import pandoc
#import OpenXML


#################################################################################33
listOfBlocks, data, godzina, hostUp, liczbaHostowWSieci = main2.main()

def writeToExcel(listOfBlocks, nazwa = "sample"):
    wb = openpyxl.Workbook()
    ws = wb.active
    ws['A1'] = "Adres IP Hosta"
    ws['B1'] = "Opoźnienie"
    ws['C1'] = "Nazwa"
    ws['D1'] = "Numer Portu"
    ws['E1'] = "Typ Uslugi"
    ws['F1'] = "Filtrowany/Otwary"
    ws['G1'] = "Nazwa Usługi"
    ws['H1'] = "System Operacyjny"
    komorka = 2
    for i in listOfBlocks:
        ws['A'+str(komorka)] = str(i.ip).split(',')[0]
        ws['B' + str(komorka)] = str(i.latency).split(',')[0]
        ws['C'+ str(komorka)] =  str(i.name).split(',')[0]
        OS  = str(i.system).replace(", ', '']", "")
        OS = str(OS).replace(",", "\n")
        ws['H' + str(komorka)] = str(OS)
        for j in str(i.portsList).split('), ('):
            try:
                j = j.replace("[('", "")
                j = j.replace(")]", "")
                j = j.replace("', '", " ")
                j = j.replace("'", "")
                j = j.replace(",", "")
                ws['D' + str(komorka)] = j.split(' ')[0]
                ws['E' + str(komorka)] = j.split(' ')[1]
                ws['F' + str(komorka)] = j.split(' ')[2]
                ws['G' + str(komorka)] = j.split(' ')[3]
                komorka+=1
            except IndexError:
                ws['D' + str(komorka)] = ''

    wb.save(nazwa+'.xlsx')
    return int(komorka)

def createTabel(doc, rows, columns):
    table = doc.add_table(rows=5, cols=3)  # x, y
    column = table.columns
    for k in column:
        k.width = Cm(3)  # szerokosc utworzonej
    return table


def fillTabel(doc, table, rows, columns, width=13):
    column = table.columns
    for k in column:
        k.width = Cm(width) #szerokosc utworzonej

def chooseStyle(doc, text="probaParagraph"):
    my_styles = docx.Document().styles
    p_style = my_styles.add_style('StylBlueEnergy', WD_STYLE_TYPE.PARAGRAPH)

    p_style.base_style = my_styles['Normal']
    p_style.paragraph_format.space_before = Pt(0)
    p_style.paragraph_format.space_after = Pt(10)

    #ch_style = my_styles.add_style('CzcionkaBlueEnergy', WD_STYLE_TYPE.CHARACTER)
    #ch_style.base_style = my_styles['Default Paragraph Font']
    p_style.font.name = 'Century Gothic'
    p_style.font.size = Pt(12)

    new_paragraph = doc.add_paragraph(text )
    new_paragraph.text = text

    return p_style


def writeToWord(rows, columns): #tabeleczka do raportu
    doc = docx.Document()
    style = doc.styles['Normal']
    style.hidden = False
    style.quick_style = True
    style.priority = 1
    p_style = chooseStyle(doc)

    paragraph = doc.add_paragraph('Hosty Uslugi')
    paragraph.style = p_style
    paragraph.text = "fewrewrwe"


    table = createTabel(doc, rows, columns)
    fillTabel(doc, table, rows, columns)

    #cell = table.cell(0,1) #wiersz, kolumna
    doc.save('nowyTabela.docx')

def main():
    #komorka = writeToExcel()
    writeToWord(5, 3)
    #writeToWriter()

if __name__ == '__main__':
    main()